from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Length, Pt, RGBColor
from docx.text.run import Run
from PIL import Image

from .models import ContentBlock, DocumentIR

CJK_FONT = "Arial Unicode MS"


def _shade_run(run: Run, fill: str = "FFF2CC") -> None:
    element = run._r
    properties = element.get_or_add_rPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def _set_run_font(run: Run, size: Length | None = None) -> None:
    run.font.name = CJK_FONT
    if size is not None:
        run.font.size = size
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), CJK_FONT)
    fonts.set(qn("w:hAnsi"), CJK_FONT)
    fonts.set(qn("w:eastAsia"), CJK_FONT)


def _write_text_block(document: DocumentType, block: ContentBlock, mark_review: bool) -> None:
    display_text = block.text
    if block.type == "heading":
        level = min(max(int(block.style.get("heading_level", 1)), 1), 3)
        paragraph = document.add_heading(level=level)
    elif block.type == "list":
        ordered = bool(re.match(r"^\s*\d+[.)、]", block.text))
        style = "List Number" if ordered else "List Bullet"
        display_text = (
            re.sub(r"^\s*\d+[.)、]\s*", "", block.text)
            if ordered
            else re.sub(r"^\s*[•·\-–]\s*", "", block.text)
        )
        paragraph = document.add_paragraph(style=style)
    else:
        paragraph = document.add_paragraph()
    run = paragraph.add_run(display_text)
    _set_run_font(run)
    run.bold = bool(block.style.get("bold", False))
    run.italic = bool(block.style.get("italic", False))
    if mark_review and block.needs_review:
        _shade_run(run)
        note = paragraph.add_run("  【需要人工核对】")
        _set_run_font(note)
        note.bold = True
        note.font.color.rgb = RGBColor(160, 95, 0)


def _write_table(document: DocumentType, block: ContentBlock, mark_review: bool) -> None:
    rows = block.table_rows
    if not rows:
        _write_text_block(document, block, mark_review)
        return
    width = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for column_index in range(width):
            value = row[column_index] if column_index < len(row) else ""
            cell = table.cell(row_index, column_index)
            cell.text = value
            for run in cell.paragraphs[0].runs:
                _set_run_font(run)
            if mark_review and block.needs_review:
                for run in cell.paragraphs[0].runs:
                    _shade_run(run)


def _write_image(document: DocumentType, block: ContentBlock) -> None:
    if not block.image_path or not Path(block.image_path).exists():
        return
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    try:
        with Image.open(block.image_path) as source:
            pixel_width, pixel_height = source.size
            dpi_value = source.info.get("dpi", (96.0, 96.0))
        dpi_x = float(dpi_value[0]) if isinstance(dpi_value, tuple) else 96.0
        dpi_y = float(dpi_value[1]) if isinstance(dpi_value, tuple) else 96.0
        dpi_x = dpi_x if 36 <= dpi_x <= 1200 else 96.0
        dpi_y = dpi_y if 36 <= dpi_y <= 1200 else 96.0
        natural_width = pixel_width / dpi_x
        natural_height = pixel_height / dpi_y
        max_width, max_height = 5.8, 7.0
        scale = min(max_width / natural_width, max_height / natural_height, 1.0)
        width = max(natural_width * scale, 0.1)
        height = max(natural_height * scale, 0.1)
        run.add_picture(block.image_path, width=Inches(width), height=Inches(height))
    except Exception:
        paragraph.add_run("[图片无法嵌入，请查看原 PDF]")


def write_docx(document_ir: DocumentIR, output: Path, mark_review: bool = True) -> Path:
    output.parent.mkdir(parents=True, exist_ok=True)
    temporary = output.with_suffix(output.suffix + ".partial")
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = document.styles["Normal"]
    normal.font.name = CJK_FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), CJK_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CJK_FONT)
    for style_name in (
        "Title",
        "Heading 1",
        "Heading 2",
        "Heading 3",
        "Caption",
        "List Bullet",
        "List Number",
    ):
        style = document.styles[style_name]
        style.font.name = CJK_FONT
        style_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        style_fonts.set(qn("w:ascii"), CJK_FONT)
        style_fonts.set(qn("w:hAnsi"), CJK_FONT)
        style_fonts.set(qn("w:eastAsia"), CJK_FONT)

    for page_index, page in enumerate(document_ir.pages):
        if page_index:
            document.add_section(WD_SECTION.NEW_PAGE)
        for block in page.blocks:
            if block.type in {"header", "footer"}:
                continue
            if block.type == "table":
                _write_table(document, block, mark_review)
            elif block.type == "image":
                _write_image(document, block)
            else:
                _write_text_block(document, block, mark_review)
        if not page.blocks:
            paragraph = document.add_paragraph("[原 PDF 此页未提取到可用内容，需要人工核对]")
            _shade_run(paragraph.runs[0])

    if document_ir.issues:
        document.add_page_break()
        document.add_heading("需要人工核对的问题清单", level=1)
        for issue in document_ir.issues:
            bbox = issue.bbox
            coordinate = (
                f"({bbox.x0:.1f}, {bbox.y0:.1f}, {bbox.x1:.1f}, {bbox.y1:.1f})"
                if bbox
                else "整页"
            )
            paragraph = document.add_paragraph(style="List Bullet")
            run = paragraph.add_run(
                f"第 {issue.page} 页 · {issue.issue_type} · {issue.message} · 坐标 {coordinate} · "
                f"置信度 {issue.confidence:.1%}"
            )
            _set_run_font(run)
            _shade_run(run)
    document.save(str(temporary))
    temporary.replace(output)
    return output
